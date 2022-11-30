from sqlalchemy.orm import Session

from ..database import Schemas

from fastapi import UploadFile
import shutil
import docx

from . import Models


def get_collections(db: Session):
    collections= db.query(Models.Collections).all()    
    return collections


def get_questions(db: Session, collections_id: int):
    questions = db.query(Models.Questions).filter(Models.Questions.collections_id == collections_id).all()
    
    data = []

    for question in questions:
        data.append({
            'id': question.id,
            'title': question.title,
            'answers': [question.answer1, question.answer2, question.answer3, question.answer4]
        })
    
    data = sorted(data, key=lambda k: k['id'])    

    return (data != []) and data or "No questions"


def create_questions(db: Session, file: UploadFile):
    dir = f'./app/data/{file.filename}'
    with open(dir, 'wb') as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    doc = docx.Document(dir)

    data = {'title': doc.paragraphs[0].text,'questions': []}

    for i in range (0, doc.paragraphs.__len__()):
        dic = {}
        if doc.paragraphs[i].text.startswith('Câu'):
            dic["question"] = doc.paragraphs[i].text
            dic["answer"] = [doc.paragraphs[i+1].text, doc.paragraphs[i+2].text, doc.paragraphs[i+3].text, doc.paragraphs[i+4].text]
            dic["correct_answer"] = doc.paragraphs[i+5].text
            data['questions'].append(dic)

    db_collection = Models.Collections(
        title = data['title']
    )
    db.add(db_collection)
    db.commit()
    db.refresh(db_collection)
    
    for i in range (0, data['questions'].__len__()):

        correct_answer = ""
        
        if data['questions'][i]['correct_answer'].replace(" ", "").replace("\n", "").replace("\t", "")[-1:] == "A":
            correct_answer = data['questions'][i]['answer'][0]
        elif data['questions'][i]['correct_answer'].replace(" ", "").replace("\n", "").replace("\t", "")[-1:] == "B":
            correct_answer = data['questions'][i]['answer'][1]
        elif data['questions'][i]['correct_answer'].replace(" ", "").replace("\n", "").replace("\t", "")[-1:] == "C":
            correct_answer = data['questions'][i]['answer'][2]
        elif data['questions'][i]['correct_answer'].replace(" ", "").replace("\n", "").replace("\t", "")[-1:] == "D":
            correct_answer = data['questions'][i]['answer'][3]
        else:
            correct_answer = "None"

        db_question = Models.Questions(
            title = data['questions'][i]['question'],
            answer1 = data['questions'][i]['answer'][0],
            answer2 = data['questions'][i]['answer'][1],
            answer3 = data['questions'][i]['answer'][2],
            answer4 = data['questions'][i]['answer'][3],
            correct_answer = correct_answer,
            collections_id = db_collection.id
        )

        db.add(db_question)
        db.commit()
        db.refresh(db_question)

    return data


def results(db: Session, test: Schemas.Test):
    data = db.query(Models.Questions).filter(Models.Questions.collections_id == test.collections_id).all()
    correct_answer = []
    for question in data:
        item = {}
        item['id'] = question.id
        item['title'] = question.title
        item['correct_answer'] = question.correct_answer
        correct_answer.append(item)
    
    score = 0
    wrong_answer = []

    for i in test.answers:
        for j in correct_answer:
            if i['id'] == j['id']:
                if i['answer'] == j['correct_answer']:
                    score += 1
                else:
                    wrong_answer.append({"id": i['id'], "question": j['title'] , "answer": i['answer'], "correct_answer": j['correct_answer']})
                    
    return {"score": score, "wrong_answer": wrong_answer}

def delete_collection(db: Session, collections_id: int):
    db.query(Models.Questions).filter(Models.Questions.collections_id == collections_id).delete()
    db.commit()
    db.query(Models.Collections).filter(Models.Collections.id == collections_id).delete()
    db.commit()
    return True